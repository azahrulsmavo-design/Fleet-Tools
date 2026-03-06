import pypdf
import docx

print("--- DOCX Text ---")
try:
    doc = docx.Document('Template.docx')
    for p in doc.paragraphs:
        if p.text.strip():
            print(p.text.strip())
    
    print("\n--- Tables in DOCX ---")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    print("TABLE CELL:", cell.text.strip())
except Exception as e:
    print("Error reading DOCX:", e)

print("\n--- PDF Fields ---")
try:
    reader = pypdf.PdfReader('Template.pdf')
    fields = reader.get_fields()
    if fields:
        for k, v in fields.items():
            print(k, v.get('/V', ''))
    else:
        print("No fields found in PDF.")
except Exception as e:
    print("Error reading PDF:", e)
